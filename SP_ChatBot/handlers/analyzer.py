import io
import os
import re
import tempfile
from aiogram import Dispatcher, F
from aiogram.fsm.context import FSMContext
from aiogram.types import CallbackQuery, Message, InlineKeyboardMarkup, InlineKeyboardButton, \
    BufferedInputFile
from docx import Document as DocxDocument
from sqlalchemy import select
from db.session import SessionLocal
from handlers.checking_criteria import contains_label0_features, contains_label1_features
from keyboards.inline import get_analysis_menu_keyboard, get_back_to_analysis_menu_keyboard
from models.bases import InterviewAnalysisResult
from states.fsm import InterviewAnalysis
from transformers import pipeline

analyzer_pipeline = pipeline("text-classification", model="ekas03/my_genspec_analyzer")

def analyzer(dp: Dispatcher):
    @dp.callback_query(F.data == "analysis")
    async def show_analysis_menu(callback: CallbackQuery):
        await callback.message.edit_text(
            "Выберите действие:",
            reply_markup=get_analysis_menu_keyboard()
        )
        await callback.answer()

    @dp.callback_query(F.data == "my_interviews")
    async def show_user_interviews(callback: CallbackQuery):
        user_id = callback.from_user.id
        async with SessionLocal() as session:
            result = await session.execute(
                select(InterviewAnalysisResult).where(InterviewAnalysisResult.telegram_id == user_id)
            )
            interviews = result.scalars().all()

        if not interviews:
            await callback.message.edit_text("У вас пока нет сохранённых интервью!", reply_markup=get_back_to_analysis_menu_keyboard())
            await callback.answer()
            return

        keyboard = InlineKeyboardMarkup(
            inline_keyboard=[[InlineKeyboardButton(text=i.title, callback_data=f"interview_{i.id}")]
                             for i in interviews] + [[InlineKeyboardButton(text="🔙 Назад", callback_data="analysis")]]
        )

        await callback.message.edit_text("Ваши интервью:", reply_markup=keyboard)
        await callback.answer()

    @dp.callback_query(F.data.startswith("interview_"))
    async def show_interview_details(callback: CallbackQuery):
        interview_id = int(callback.data.split("_")[1])

        async with SessionLocal() as session:
            interview = await session.get(InterviewAnalysisResult, interview_id)

        if not interview:
            await callback.message.answer("Интервью не найдено.")
            return

        msg = (
            f"📊 <b>{interview.title}</b>\n"
            f"Всего предложений: {interview.total_sentences}\n"
            f"ОБЩИХ: {interview.general_count}\n"
            f"ЧАСТНЫХ: {interview.specific_count}\n\n"
            f"В результате {interview.majority}"
        )

        keyboard = InlineKeyboardMarkup(
            inline_keyboard=[
                [InlineKeyboardButton(text="📤 Скачать результат", callback_data=f"download_{interview_id}")],
                [InlineKeyboardButton(text="🔙 Назад", callback_data="my_interviews")]
            ]
        )

        await callback.message.edit_text(msg, reply_markup=keyboard, parse_mode="HTML")
        await callback.answer()

        @dp.callback_query(F.data.startswith("download_"))
        async def send_interview_docx(callback: CallbackQuery):
            interview_id = int(callback.data.split("_")[1])

            async with SessionLocal() as session:
                interview = await session.get(InterviewAnalysisResult, interview_id)

            if not interview:
                await callback.message.answer("Интервью не найдено.")
                return

            per_general = (interview.general_count / interview.total_sentences) * 100 if interview.total_sentences > 0 else 0
            per_specific = (interview.specific_count / interview.total_sentences) * 100 if interview.total_sentences > 0 else 0

            doc = DocxDocument()
            doc.add_heading(interview.title, 0)
            doc.add_paragraph(f"Всего предложений: {interview.total_sentences}")
            doc.add_paragraph(f"ОБЩИХ: {interview.general_count}, то есть {per_general:.1f}%\n\n")
            doc.add_paragraph(f"ЧАСТНЫХ: {interview.specific_count}, то есть {per_specific:.1f}%\n\n")
            doc.add_paragraph(f"В результате {interview.majority}")
            doc.add_heading("Полный текст интервью:", level=1)
            doc.add_paragraph(interview.full_text)
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)

            file = BufferedInputFile(buf.read(), filename=f"{interview.title}.docx")
            await callback.message.answer_document(file)
            await callback.answer()

    @dp.callback_query(F.data == "analyze_interview")
    async def start_analysis(callback: CallbackQuery, state: FSMContext):
        await callback.message.edit_text(
            "Введите название интервью:",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🔙 Назад", callback_data="analysis")]]))
        await state.set_state(InterviewAnalysis.waiting_for_title)
        await callback.answer()

    @dp.message(InterviewAnalysis.waiting_for_title)
    async def get_interview_title(message: Message, state: FSMContext):
        await state.update_data(title=message.text)
        await message.answer(
            "Пожалуйста, отправьте файл с расширением .docx, где вопросы начинаются с \"В:\", а ответы с \"О:\"",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🔙 Назад", callback_data="analyze_interview")]])
        )
        await state.set_state(InterviewAnalysis.waiting_for_docx)

    @dp.message(InterviewAnalysis.waiting_for_docx, F.document)
    async def get_docx_file(message: Message, state: FSMContext):
        doc = message.document
        if not doc.file_name.endswith(".docx"):
            await message.answer("Пожалуйста, отправьте файл с расширением .docx, где вопросы начинаются с \"В:\", а ответы с \"О:\"")
            return

        with tempfile.TemporaryDirectory() as tmpdir:
            file_path = os.path.join(tmpdir, doc.file_name)
            await message.bot.download(doc, destination=file_path)

            docx = DocxDocument(file_path)
            full_text = "\n".join([para.text for para in docx.paragraphs])

        pattern = re.compile(r"О:\s*(.*?)(?=\s*В:|\s*$)", re.DOTALL)
        matches = pattern.findall(full_text)
        if not matches:
            await message.answer("Не удалось найти фрагменты между 'О:' и 'В:'.")
            return

        interview_text = "\n".join(matches)

        data = await state.get_data()
        title = data.get("title")

        try:
            msg, _ = await analyze_and_save(interview_text, message.from_user.id, title)
        except Exception as e:
            await message.answer(str(e))
            await state.clear()
            return

        await message.answer(msg, reply_markup=get_back_to_analysis_menu_keyboard())
        await state.clear()

def split_into_sentences(text: str) -> list[str]:
    sentences = re.split(r'(?<=[.!?])[\s\n]+', text.strip())
    return [s.strip() for s in sentences if s.strip()]

async def analyze_and_save(text: str, user_id: int, title: str) -> tuple[str, InterviewAnalysisResult]:
        sentences = split_into_sentences(text)
        if not sentences:
            raise ValueError("Невозможно разделить текст на предложения.")

        try:
            results = analyzer_pipeline(sentences, truncation=True)
        except Exception as e:
            raise RuntimeError(f"Ошибка при анализе модели: {e}")

        grammar_label0_count = sum(1 for s in sentences if contains_label0_features(s))
        grammar_label1_count = sum(1 for s in sentences if contains_label1_features(s))

        general_count = sum(1 for r in results if r["label"].lower() == "label_0") + grammar_label0_count
        specific_count = sum(1 for r in results if r["label"].lower() == "label_1") + grammar_label1_count
        total = len(results) + grammar_label0_count + grammar_label1_count

        majority = (
            "больше ОБЩИХ признаков, а это значит, что человек в первую очередь видит все процессы в целом, фокусируется на сути, концепции и стратегии!" if general_count > specific_count else
            "больше ЧАСТНЫХ признаков, а это значит, что человек в первую очередь замечает детали, ориентируется на конкретику и пошаговое решение поставленной задачи!" if specific_count > general_count else
            "равное количество ОБЩИХ и ЧАСТНЫХ признаков. Это может говорить и о том, что в человеке есть баланс \"общих\" и \"частных\" признаков, но и о том, что отправленных на анализ данных недостаточно, поэтому необходимо повторно провести интервью и отправить его на цифровой анализ!"
        )

        per_general = (general_count / total) * 100 if total > 0 else 0
        per_specific = (specific_count / total) * 100 if total > 0 else 0

        result_text = (
            f"📊 Результаты анализа интервью:\n\n"
            f"Всего выявлено предложений: {total}\n\n"
            f"Количество предложений, отмеченных как GENERAL (ОБЩЕЕ): {general_count}, то есть {per_general:.1f}%\n\n"
            f"Количество предложений, отмеченных как SPECIFIC (ЧАСТНОЕ): {specific_count}, то есть {per_specific:.1f}%)\n\n"
            f"В результате анализа было найдено {majority}"
        )

        result = InterviewAnalysisResult(
            title=title,
            telegram_id=user_id,
            total_sentences=total,
            general_count=general_count,
            specific_count=specific_count,
            majority=majority,
            full_text=text
        )

        async with SessionLocal() as session:
            session.add(result)
            await session.commit()

        return result_text, result